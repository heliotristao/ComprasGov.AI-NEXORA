import json
import io
import os
import openai
from docx import Document
from PyPDF2 import PdfReader
from sentence_transformers import SentenceTransformer
from sqlalchemy.sql import text

from app.models import Artifact, ArtifactIndex
from app.db.session import SessionLocal

# --- Mock S3 Content ---
# In a real implementation, this would be replaced with actual S3 calls
MOCK_S3_STORE = {
    "test/key.json": io.BytesIO(json.dumps({"content": "This is a test JSON file."}).encode()),
    "test/key.docx": io.BytesIO(),  # Will be populated by a mock docx
    "test/key.pdf": io.BytesIO(),  # Will be populated by a mock pdf
}
# Create a mock docx for testing
docx_doc = Document()
docx_doc.add_paragraph("This is a test DOCX file.")
docx_doc.save(MOCK_S3_STORE["test/key.docx"])
MOCK_S3_STORE["test/key.docx"].seek(0)


def extract_text_from_artifact(artifact: Artifact) -> str:
    """Extracts text from an artifact."""
    file_content = MOCK_S3_STORE.get(artifact.s3_key)
    if not file_content:
        return ""

    if artifact.filename.endswith(".json"):
        data = json.load(file_content)
        return json.dumps(data)
    elif artifact.filename.endswith(".docx"):
        doc = Document(file_content)
        return " ".join([p.text for p in doc.paragraphs])
    elif artifact.filename.endswith(".pdf"):
        reader = PdfReader(file_content)
        return " ".join([page.extract_text() for page in reader.pages])
    return ""

def generate_embedding(text: str) -> list[float]:
    """Generates an embedding for the given text."""
    try:
        # Try OpenAI first
        openai.api_key = os.getenv("OPENAI_API_KEY", "placeholder")
        response = openai.Embedding.create(input=text, model="text-embedding-ada-002")
        return response["data"][0]["embedding"]
    except Exception:
        # Fallback to Sentence-Transformers
        model = SentenceTransformer("all-MiniLM-L6-v2")
        return model.encode(text).tolist()

def _index_artifact_with_db(db, artifact_id: str):
    """Helper function to index an artifact with a given db session."""
    artifact = db.query(Artifact).filter(Artifact.id == artifact_id).first()
    if not artifact:
        return

    text = extract_text_from_artifact(artifact)
    embedding = generate_embedding(text)
    summary = text[:500]

    artifact_index = ArtifactIndex(
        artifact_id=artifact.id,
        doc_type=artifact.doc_type.value,
        org_id=artifact.org_id,
        embedding=embedding,
        summary=summary,
    )
    db.add(artifact_index)
    db.commit()


def index_artifact(artifact_id: str):
    """Indexes an artifact by extracting its text, generating an embedding, and saving it to the database."""
    db = SessionLocal()
    try:
        _index_artifact_with_db(db, artifact_id)
    finally:
        db.close()

def search_artifacts(query: str, semantic: bool = True) -> list:
    """Searches for artifacts using a hybrid search approach."""
    db = SessionLocal()
    try:
        # Full-text search
        fts_query = text("SELECT artifact_id, summary FROM artifact_index WHERE to_tsvector('english', summary) @@ to_tsquery('english', :query)")
        fts_results = db.execute(fts_query, {"query": query}).fetchall()

        if not semantic:
            return [{"id": str(r[0]), "summary": r[1], "score": 1.0} for r in fts_results]

        # Semantic search
        query_embedding = generate_embedding(query)
        # The l2_distance in pgvector is the squared Euclidean distance.
        # Cosine similarity can be calculated from it, but for simplicity, we use l2_distance directly.
        # A lower distance means a better match.
        semantic_query = text("SELECT artifact_id, summary, embedding <-> :embedding AS distance FROM artifact_index ORDER BY distance LIMIT 10")
        semantic_results = db.execute(semantic_query, {"embedding": str(query_embedding)}).fetchall()

        # Hybrid scoring (simple approach)
        results = {}
        for r in fts_results:
            results[r[0]] = {"summary": r[1], "score": 0.3} # Base score for FTS match

        for r in semantic_results:
            if r[0] in results:
                results[r[0]]["score"] += 0.7 * (1 - r[2]) # Higher score for lower distance
            else:
                results[r[0]] = {"summary": r[1], "score": 0.7 * (1 - r[2])}

        return [{"id": str(k), "summary": v["summary"], "score": v["score"]} for k, v in results.items()]

    finally:
        db.close()
