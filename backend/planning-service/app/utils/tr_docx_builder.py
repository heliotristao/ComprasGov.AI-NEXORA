"""
DOCX Builder for Termo de Referência (TR)
"""
from io import BytesIO
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.db.models.tr import TR, TRType

# A simplified template structure until the model is fully clarified.
# In a real scenario, this would come from `app.db.models.templates_gestao.py`.
from typing import Any, Dict

Template = Dict[str, Any]


def generate_docx(tr: TR, template: Template, watermark: str | None = None) -> BytesIO:
    """
    Generates a DOCX document from a TR object and a template.

    Args:
        tr: The TermoReferencia object.
        template: The template object with styling and structure info.
        watermark: Optional text to be added as a watermark.

    Returns:
        A BytesIO stream containing the generated DOCX document.
    """
    doc = Document()
    _apply_template_settings(doc, template)

    _add_header(doc, template.get("header_text", "NEXORA-ComprasGov.AI"))
    _add_title(doc, "Termo de Referência")
    _add_main_content(doc, tr)
    _add_gaps_report(doc, tr.gaps_report)

    if watermark:
        _add_watermark(doc, watermark)

    _add_footer(doc, template.get("footer_text", f"Documento gerado em {tr.updated_at or tr.created_at}"))

    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream


def _apply_template_settings(doc: Document, template: Template):
    """Applies basic document settings from the template."""
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(template.get("margin_top", 1.0))
        section.bottom_margin = Inches(template.get("margin_bottom", 1.0))
        section.left_margin = Inches(template.get("margin_left", 1.0))
        section.right_margin = Inches(template.get("margin_right", 1.0))


def _add_header(doc: Document, text: str):
    """Adds a header to the document."""
    header = doc.sections[0].header
    p = header.paragraphs[0]
    p.text = text
    p.style = 'Header'
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _add_footer(doc: Document, text: str):
    """Adds a footer to the document."""
    footer = doc.sections[0].footer
    p = footer.paragraphs[0]
    p.text = text
    p.style = 'Footer'
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _add_title(doc: Document, title: str):
    """Adds a centered title to the document."""
    doc.add_heading(title, level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()


def _add_main_content(doc: Document, tr: TR):
    """Adds the main TR content, switching based on type."""
    doc.add_heading("1. Objeto", level=2)
    doc.add_paragraph(tr.dados.get("objeto", "Não especificado."))

    if tr.type == TRType.BEM:
        _add_bem_sections(doc, tr.dados)
    else:  # SERVICO
        _add_servico_sections(doc, tr.dados)

    doc.add_heading("2. Justificativa e Fundamentação", level=2)
    doc.add_paragraph(tr.dados.get("justificativa", "Não especificado."))


def _add_bem_sections(doc: Document, data: Dict[str, Any]):
    """Adds sections specific to 'Bens'."""
    doc.add_heading("1.1. Especificações Técnicas", level=3)
    specs = data.get("especificacoes_tecnicas", "Nenhuma especificação fornecida.")
    doc.add_paragraph(specs)

    doc.add_heading("1.2. Quantitativos", level=3)
    quant = data.get("quantitativos", "Nenhum quantitativo fornecido.")
    doc.add_paragraph(str(quant))


def _add_servico_sections(doc: Document, data: Dict[str, Any]):
    """Adds sections specific to 'Serviços'."""
    doc.add_heading("1.1. Níveis de Serviço (SLAs)", level=3)
    slas = data.get("slas", "Nenhum SLA fornecido.")
    doc.add_paragraph(slas)

    doc.add_heading("1.2. Metodologia de Execução", level=3)
    method = data.get("metodologia", "Nenhuma metodologia fornecida.")
    doc.add_paragraph(method)


def _add_gaps_report(doc: Document, report: Dict[str, Any]):
    """Adds the gaps report as a final section."""
    if not report:
        return

    doc.add_page_break()
    doc.add_heading("Anexo - Relatório de Gaps (ETP -> TR)", level=1)

    for key, value in report.items():
        doc.add_heading(key.replace("_", " ").title(), level=3)
        if isinstance(value, list):
            for item in value:
                doc.add_paragraph(str(item), style='List Bullet')
        else:
            doc.add_paragraph(str(value))


def _add_watermark(doc: Document, text: str):
    """
    A simple textual watermark implementation.
    Note: Real watermarking in python-docx is complex and might require
    modifying the underlying XML directly. This is a simplified approach.
    """
    for section in doc.sections:
        header = section.header
        p = header.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.size = Pt(48)
        run.font.bold = True
        run.font.color.rgb = (200, 200, 200)
