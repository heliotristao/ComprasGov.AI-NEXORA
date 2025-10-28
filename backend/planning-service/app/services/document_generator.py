"""
Serviço de geração de documentos DOCX e PDF para ETP e TR
"""

from typing import Dict, Any, Optional
from datetime import datetime
from io import BytesIO
import os

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from sqlalchemy.orm import Session

from app.db.models.etp_modular import DocumentoETP, DocumentoTR
from app.db.models.templates_gestao import ModeloInstitucional, Instituicao


class DocumentGenerator:
    """
    Gerador de documentos DOCX e PDF
    """
    
    def __init__(self, db: Session):
        self.db = db
    
    def gerar_etp_docx(
        self,
        documento_id: int,
        output_path: Optional[str] = None
    ) -> str:
        """
        Gera documento ETP em formato DOCX
        
        Args:
            documento_id: ID do documento ETP
            output_path: Caminho para salvar o arquivo (opcional)
        
        Returns:
            Caminho do arquivo gerado
        """
        # Buscar documento
        documento = self.db.query(DocumentoETP).filter(
            DocumentoETP.id == documento_id
        ).first()
        
        if not documento:
            raise ValueError(f"Documento {documento_id} não encontrado")
        
        # Buscar template
        template = self.db.query(ModeloInstitucional).filter(
            ModeloInstitucional.id == documento.template_id
        ).first()
        
        if not template:
            raise ValueError(f"Template {documento.template_id} não encontrado")
        
        # Buscar instituição
        instituicao = self.db.query(Instituicao).filter(
            Instituicao.id == template.instituicao_id
        ).first()
        
        # Criar documento Word
        doc = Document()
        
        # Aplicar configurações do template
        self._aplicar_configuracoes(doc, template, instituicao)
        
        # Adicionar cabeçalho
        self._adicionar_cabecalho(doc, template, instituicao)
        
        # Adicionar título
        self._adicionar_titulo(doc, "ESTUDO TÉCNICO PRELIMINAR - ETP")
        
        # Adicionar identificação
        self._adicionar_identificacao(doc, documento, template)
        
        # Adicionar conteúdo das seções
        self._adicionar_secoes(doc, documento, template)
        
        # Adicionar rodapé
        self._adicionar_rodape(doc, template, instituicao)
        
        # Salvar documento
        if not output_path:
            output_path = f"/tmp/etp_{documento_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        
        doc.save(output_path)
        
        return output_path
    
    def gerar_tr_docx(
        self,
        documento_id: int,
        output_path: Optional[str] = None
    ) -> str:
        """
        Gera documento TR em formato DOCX
        
        Args:
            documento_id: ID do documento TR
            output_path: Caminho para salvar o arquivo (opcional)
        
        Returns:
            Caminho do arquivo gerado
        """
        # Buscar documento
        documento = self.db.query(DocumentoTR).filter(
            DocumentoTR.id == documento_id
        ).first()
        
        if not documento:
            raise ValueError(f"Documento {documento_id} não encontrado")
        
        # Buscar template
        template = self.db.query(ModeloInstitucional).filter(
            ModeloInstitucional.id == documento.template_id
        ).first()
        
        if not template:
            raise ValueError(f"Template {documento.template_id} não encontrado")
        
        # Buscar instituição
        instituicao = self.db.query(Instituicao).filter(
            Instituicao.id == template.instituicao_id
        ).first()
        
        # Criar documento Word
        doc = Document()
        
        # Aplicar configurações do template
        self._aplicar_configuracoes(doc, template, instituicao)
        
        # Adicionar cabeçalho
        self._adicionar_cabecalho(doc, template, instituicao)
        
        # Adicionar título
        self._adicionar_titulo(doc, "TERMO DE REFERÊNCIA - TR")
        
        # Adicionar identificação
        self._adicionar_identificacao_tr(doc, documento, template)
        
        # Adicionar conteúdo das seções
        self._adicionar_secoes_tr(doc, documento, template)
        
        # Adicionar rodapé
        self._adicionar_rodape(doc, template, instituicao)
        
        # Salvar documento
        if not output_path:
            output_path = f"/tmp/tr_{documento_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        
        doc.save(output_path)
        
        return output_path
    
    def _aplicar_configuracoes(
        self,
        doc: Document,
        template: ModeloInstitucional,
        instituicao: Optional[Instituicao]
    ):
        """
        Aplica configurações de formatação do template
        """
        config = template.configuracao_documento or {}
        
        # Configurar margens
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(config.get("margem_superior", 1))
            section.bottom_margin = Inches(config.get("margem_inferior", 1))
            section.left_margin = Inches(config.get("margem_esquerda", 1))
            section.right_margin = Inches(config.get("margem_direita", 1))
    
    def _adicionar_cabecalho(
        self,
        doc: Document,
        template: ModeloInstitucional,
        instituicao: Optional[Instituicao]
    ):
        """
        Adiciona cabeçalho do documento
        """
        config = template.configuracao_documento or {}
        
        # Adicionar logo se disponível
        if instituicao and instituicao.logo_url:
            # TODO: Baixar e adicionar logo
            pass
        
        # Adicionar nome da instituição
        if instituicao:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(instituicao.nome.upper())
            run.bold = True
            run.font.size = Pt(14)
            
            if instituicao.sigla:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(instituicao.sigla)
                run.font.size = Pt(12)
        
        # Adicionar texto customizado do cabeçalho
        if config.get("texto_cabecalho"):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(config["texto_cabecalho"])
            run.font.size = Pt(10)
        
        # Linha separadora
        doc.add_paragraph("_" * 80)
    
    def _adicionar_titulo(self, doc: Document, titulo: str):
        """
        Adiciona título principal do documento
        """
        doc.add_paragraph()  # Espaço
        
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(titulo)
        run.bold = True
        run.font.size = Pt(16)
        
        doc.add_paragraph()  # Espaço
    
    def _adicionar_identificacao(
        self,
        doc: Document,
        documento: DocumentoETP,
        template: ModeloInstitucional
    ):
        """
        Adiciona seção de identificação do documento
        """
        p = doc.add_paragraph()
        p.add_run("IDENTIFICAÇÃO").bold = True
        
        # Número do documento
        doc.add_paragraph(f"Documento nº: {documento.id}")
        
        # Data de criação
        doc.add_paragraph(
            f"Data de Elaboração: {documento.created_at.strftime('%d/%m/%Y')}"
        )
        
        # Versão do template
        doc.add_paragraph(f"Modelo: {template.nome} (v{template.versao})")
        
        doc.add_paragraph()  # Espaço
    
    def _adicionar_secoes(
        self,
        doc: Document,
        documento: DocumentoETP,
        template: ModeloInstitucional
    ):
        """
        Adiciona todas as seções do documento
        """
        estrutura = template.estrutura
        dados = documento.dados
        
        for secao in estrutura.get("secoes", []):
            # Adicionar título da seção
            p = doc.add_paragraph()
            run = p.add_run(f"{secao['ordem']}. {secao['titulo'].upper()}")
            run.bold = True
            run.font.size = Pt(12)
            
            # Adicionar nota explicativa se houver
            if secao.get("nota_explicativa"):
                p = doc.add_paragraph()
                run = p.add_run(f"Nota: {secao['nota_explicativa']}")
                run.italic = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(128, 128, 128)
            
            # Adicionar campos da seção
            secao_dados = dados.get(secao["id"], {})
            
            for campo in secao.get("campos", []):
                campo_valor = secao_dados.get(campo["id"], "")
                
                if campo_valor:
                    # Adicionar label do campo
                    p = doc.add_paragraph()
                    run = p.add_run(f"{campo['label']}:")
                    run.bold = True
                    run.font.size = Pt(10)
                    
                    # Adicionar valor do campo
                    if campo["tipo"] in ["textarea", "text"]:
                        doc.add_paragraph(str(campo_valor))
                    elif campo["tipo"] == "radio":
                        doc.add_paragraph(f"☑ {campo_valor}")
                    elif campo["tipo"] == "select":
                        doc.add_paragraph(f"→ {campo_valor}")
                    else:
                        doc.add_paragraph(str(campo_valor))
                    
                    # Indicar se foi gerado por IA
                    if campo["id"] in documento.campos_gerados_ia:
                        p = doc.add_paragraph()
                        run = p.add_run("✨ Conteúdo gerado com assistência de IA")
                        run.italic = True
                        run.font.size = Pt(8)
                        run.font.color.rgb = RGBColor(100, 100, 200)
                        
                        # Adicionar score de confiança
                        score = documento.scores_confianca_ia.get(campo["id"], 0)
                        if score:
                            run = p.add_run(f" (Confiança: {score*100:.0f}%)")
                            run.font.size = Pt(8)
            
            doc.add_paragraph()  # Espaço entre seções
    
    def _adicionar_rodape(
        self,
        doc: Document,
        template: ModeloInstitucional,
        instituicao: Optional[Instituicao]
    ):
        """
        Adiciona rodapé do documento
        """
        config = template.configuracao_documento or {}
        
        # Linha separadora
        doc.add_paragraph("_" * 80)
        
        # Texto customizado do rodapé
        if config.get("texto_rodape"):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(config["texto_rodape"])
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(128, 128, 128)
        
        # Data de geração
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(
            f"Documento gerado em {datetime.now().strftime('%d/%m/%Y às %H:%M')}"
        )
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(128, 128, 128)
        
        # Assinatura do sistema
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("ComprasGov.AI - NEXORA")
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(128, 128, 128)
    
    def converter_para_pdf(self, docx_path: str) -> str:
        """
        Converte documento DOCX para PDF
        
        Args:
            docx_path: Caminho do arquivo DOCX
        
        Returns:
            Caminho do arquivo PDF gerado
        """
        # TODO: Implementar conversão para PDF
        # Opções:
        # 1. LibreOffice headless: soffice --headless --convert-to pdf
        # 2. python-docx-template + reportlab
        # 3. docx2pdf (Windows only)
        # 4. pypandoc
        
        pdf_path = docx_path.replace(".docx", ".pdf")
        
        # Por enquanto, apenas retornar path mockado
        return pdf_path



    def _adicionar_identificacao_tr(
        self,
        doc: Document,
        documento: DocumentoTR,
        template: ModeloInstitucional
    ):
        """
        Adiciona seção de identificação do documento TR
        """
        p = doc.add_paragraph()
        p.add_run("IDENTIFICAÇÃO").bold = True
        
        # Número do documento
        doc.add_paragraph(f"Documento nº: {documento.id}")
        
        # Referência ao ETP
        if documento.etp_id:
            doc.add_paragraph(f"Baseado no ETP nº: {documento.etp_id}")
        
        # Data de criação
        doc.add_paragraph(
            f"Data de Elaboração: {documento.created_at.strftime('%d/%m/%Y')}"
        )
        
        # Versão do template
        doc.add_paragraph(f"Modelo: {template.nome} (v{template.versao})")
        
        doc.add_paragraph()  # Espaço
    
    def _adicionar_secoes_tr(
        self,
        doc: Document,
        documento: DocumentoTR,
        template: ModeloInstitucional
    ):
        """
        Adiciona todas as seções do documento TR
        """
        estrutura = template.estrutura
        dados = documento.dados
        
        for secao in estrutura.get("secoes", []):
            # Adicionar título da seção
            p = doc.add_paragraph()
            run = p.add_run(f"{secao['ordem']}. {secao['titulo'].upper()}")
            run.bold = True
            run.font.size = Pt(12)
            
            # Adicionar nota explicativa se houver
            if secao.get("nota_explicativa"):
                p = doc.add_paragraph()
                run = p.add_run(f"Nota: {secao['nota_explicativa']}")
                run.italic = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(128, 128, 128)
            
            # Adicionar campos da seção
            secao_dados = dados.get(secao["id"], {})
            
            for campo in secao.get("campos", []):
                campo_valor = secao_dados.get(campo["id"], "")
                
                if campo_valor:
                    # Adicionar label do campo
                    p = doc.add_paragraph()
                    run = p.add_run(f"{campo['label']}:")
                    run.bold = True
                    run.font.size = Pt(10)
                    
                    # Adicionar valor do campo
                    if campo["tipo"] in ["textarea", "text"]:
                        doc.add_paragraph(str(campo_valor))
                    elif campo["tipo"] == "radio":
                        doc.add_paragraph(f"☑ {campo_valor}")
                    elif campo["tipo"] == "select":
                        doc.add_paragraph(f"→ {campo_valor}")
                    else:
                        doc.add_paragraph(str(campo_valor))
                    
                    # Indicar se foi gerado por IA
                    if campo["id"] in documento.campos_gerados_ia:
                        p = doc.add_paragraph()
                        run = p.add_run("✨ Conteúdo gerado com assistência de IA")
                        run.italic = True
                        run.font.size = Pt(8)
                        run.font.color.rgb = RGBColor(100, 100, 200)
                        
                        # Adicionar score de confiança
                        score = documento.scores_confianca_ia.get(campo["id"], 0)
                        if score:
                            run = p.add_run(f" (Confiança: {score*100:.0f}%)")
                            run.font.size = Pt(8)
            
            doc.add_paragraph()  # Espaço entre seções

