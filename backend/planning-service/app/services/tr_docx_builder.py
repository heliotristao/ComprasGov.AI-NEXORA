import os
import tempfile
from docx import Document
from app.db.models.tr import TR

class TRDocxBuilder:
    def __init__(self, tr: TR, template_path: str):
        self.tr = tr
        if not os.path.exists(template_path):
            raise FileNotFoundError(f"Template not found at path: {template_path}")
        self.document = Document(template_path)
        self.context = self.tr.data

    def build(self) -> str:
        """
        Builds a DOCX document from a TR object using a template.
        """
        self._replace_placeholders()
        self._add_gap_report()
        self._add_watermark()
        return self._save_document()

    def _replace_placeholders(self):
        """
        Replaces placeholders in the format {{key}} with data from the TR.
        """
        context = {**self.tr.data, "objeto": self.tr.title}

        # Replace in paragraphs
        for p in self.document.paragraphs:
            text = p.text
            for key, value in context.items():
                placeholder = f"{{{{{key}}}}}"
                if placeholder in text:
                    text = text.replace(placeholder, str(value))
            if p.text != text:
                p.clear()
                p.add_run(text)

        # Replace in tables
        for table in self.document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        text = p.text
                        for key, value in context.items():
                            placeholder = f"{{{{{key}}}}}"
                            if placeholder in text:
                                text = text.replace(placeholder, str(value))
                        if p.text != text:
                            p.clear()
                            p.add_run(text)

    def _add_gap_report(self):
        """
        Adds a gap report to the document.
        """
        self.document.add_heading("Relatório de Gaps", level=2)
        if self.tr.gaps:
            for key, value in self.tr.gaps.items():
                self.document.add_paragraph(f"- {key}: {value}")
        else:
            self.document.add_paragraph("Nenhum gap encontrado.")

    def _add_watermark(self):
        """
        Adds a watermark to the document.
        """
        # Placeholder for watermark logic
        pass

    def _save_document(self) -> str:
        """
        Saves the document to a temporary file and returns the path.
        """
        temp_dir = tempfile.gettempdir()
        file_path = os.path.join(temp_dir, f"tr_{self.tr.id}.docx")
        self.document.save(file_path)
        return file_path
